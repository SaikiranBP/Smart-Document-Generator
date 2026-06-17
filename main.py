from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

# =====================================================
# DOCUMENT CONFIGURATION & BORDERS
# =====================================================

def configure_borders_not_to_surround_header_footer(doc):
    """
    Configures the document-level settings so that page borders do not
    surround the header and footer, keeping them outside the page border.
    """
    settings = doc.settings._element    
    for setting_name in ['bordersDoNotSurroundHeader', 'bordersDoNotSurroundFooter']:
        element = settings.find(qn(f'w:{setting_name}'))
        if element is None:
            element = OxmlElement(f'w:{setting_name}')
            settings.append(element)
        element.set(qn('w:val'), '1')

def add_page_border(section):
    """
    Adds a page border to the given section, offset from the text margins.
    """
    sectPr = section._sectPr

    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'text')  # Align with text margins

    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')  # Border size (thickness)
        border.set(qn('w:space'), '24')  # Spacing in points
        border.set(qn('w:color'), '000000')  # Black color
        pgBorders.append(border)

    sectPr.append(pgBorders)

# =====================================================
# PAGE NUMBERING
# =====================================================

def restart_page_numbering(section):
    """
    Sets the page numbering to restart at 1 for the given section.
    """
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:start'), '1')

def add_page_number(paragraph):
    """
    Adds a dynamic page number field to the paragraph, formatted in Times New Roman.
    """
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

# =====================================================
# TEXT FORMATTING HELPERS
# =====================================================

def add_main_heading(doc, text):
    """
    Adds a main heading (Heading 1) center-aligned, 18pt, bold, Times New Roman, 1.5 line spacing.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    return p

def add_sub_heading(doc, text):
    """
    Adds a subheading, justified, 14pt, bold, Times New Roman, 1.5 line spacing.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    return p

def add_body_text(doc, text):
    """
    Adds body text, justified, 12pt, Times New Roman, 1.5 line spacing.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p

# =====================================================
# MAIN EXECUTION FLOW
# =====================================================

def main():
    print("\n=============================================")
    print("      SMART WORD DOCUMENT GENERATOR          ")
    print("=============================================\n")

    # Ask user for general report details
    filename_input = input("Enter Filename (or press Enter for default): ").strip()
    if not filename_input:
        filename_input = "Smart Word Report"

    header_text = input("Enter Header Text: ").strip()
    footer_text = input("Enter Footer Text: ").strip()

    # Ask for number of content pages (starting on page 3)
    while True:
        try:
            num_pages = int(input("Enter the number of pages of content to create: "))
            if num_pages > 0:
                break
            print("Please enter a number greater than 0.")
        except ValueError:
            print("Invalid input. Please enter a valid integer.")

    # Collect page-specific details
    pages = []
    for i in range(1, num_pages + 1):
        print(f"\n--- Page {i} Content (will start on Document Page {i+2}) ---")
        heading = input(f"Enter Heading for Page {i}: ").strip()
        subheading = input(f"Enter Subheading for Page {i}: ").strip()
        body = input(f"Enter Body Text for Page {i}: ").strip()
        pages.append({
            "heading": heading,
            "subheading": subheading,
            "body": body
        })

    # Create Document
    doc = Document()
    
    # Configure borders not to surround header/footer
    configure_borders_not_to_surround_header_footer(doc)

    # Configure default style line spacing for the whole document (1.5 line spacing)
    style = doc.styles['Normal']
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.line_spacing = 1.5

    # -------------------------------------------------
    # Section 1: First 2 Pages (Empty, border-free)
    # -------------------------------------------------
    section1 = doc.sections[0]
    section1.top_margin = Cm(2)
    section1.bottom_margin = Cm(2)
    section1.left_margin = Cm(2)
    section1.right_margin = Cm(2)

    # Add an empty paragraph on page 1 and move to page 2
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p1.paragraph_format.line_spacing = 1.5
    doc.add_page_break()

    # Add an empty paragraph on page 2
    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p2.paragraph_format.line_spacing = 1.5

    # -------------------------------------------------
    # Section 2: Page 3 and Onwards (Content pages)
    # -------------------------------------------------
    # Create section break starting on the next page (Document Page 3)
    sec2 = doc.add_section(WD_SECTION_START.NEW_PAGE)
    
    # Unlink headers and footers from Section 1 to keep pages 1-2 blank
    sec2.header.is_linked_to_previous = False
    sec2.footer.is_linked_to_previous = False

    sec2.top_margin = Cm(2)
    sec2.bottom_margin = Cm(2)
    sec2.left_margin = Cm(2)
    sec2.right_margin = Cm(2)
    
    # Position header/footer inside the margin area, so they sit outside the border
    sec2.header_distance = Cm(0.6)
    sec2.footer_distance = Cm(0.5)

    # Add page borders to Section 2
    add_page_border(sec2)

    # Restart numbering at 1 for Section 2
    restart_page_numbering(sec2)

    # Configure Header for Section 2
    header = sec2.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    header_para.paragraph_format.line_spacing = 1.5
    header_run = header_para.add_run(header_text)
    header_run.font.name = "Times New Roman"
    header_run.font.size = Pt(10)

    # Configure Footer for Section 2 (with bottom-right page number)
    footer = sec2.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    footer_para.paragraph_format.line_spacing = 1.5
    
    # Add footer text run
    footer_run = footer_para.add_run(footer_text)
    footer_run.font.name = "Times New Roman"
    footer_run.font.size = Pt(10)

    # Add tab stop at the right margin edge to align page number to the right
    usable_width = sec2.page_width - sec2.left_margin - sec2.right_margin
    footer_para.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
    
    # Add tab space to move cursor to the right
    footer_para.add_run("\t\t")
    
    # Add page number
    add_page_number(footer_para)

    # -------------------------------------------------
    # Add content to Document (Section 2)
    # -------------------------------------------------

    for idx, page in enumerate(pages):
        # Add heading (Centered, size 18)
        if page["heading"]:
            add_main_heading(doc, page["heading"])
        
        # Add subheading (Justified, size 14)
        if page["subheading"]:
            add_sub_heading(doc, page["subheading"])
        
        # Add body text (Justified, size 12)
        if page["body"]:
            add_body_text(doc, page["body"])

        # Add page break if it's not the last page
        if idx < len(pages) - 1:
            doc.add_page_break()

    # Save the file (handling file locking and overwriting/overlapping changes gracefully)
    filename = filename_input.replace(" ", "_")
    if not filename.endswith(".docx"):
        filename += ".docx"
    while True:
        try:
            doc.save(filename)
            break
        except PermissionError:
            print(f"\n[ERROR] Permission Denied: Cannot write to '{filename}'.")
            print("The file might be open in Microsoft Word or another application.")
            response = input("Please close the document and press Enter to retry, or enter a new filename: ").strip()
            if response:
                if not response.endswith(".docx"):
                    response += ".docx"
                filename = response
    
    print("\n=============================================")
    print("Document generated successfully!")
    print(f"Saved as: {filename}")
    print("=============================================\n")

if __name__ == "__main__":
    main()