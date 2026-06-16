import streamlit as st
import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from google import genai
import os

# Import document layout helpers from main.py
from main import (
    configure_borders_not_to_surround_header_footer,
    add_page_border,
    restart_page_numbering,
    add_page_number,
    add_main_heading,
    add_sub_heading,
    add_body_text
)

# Set up page configurations
st.set_page_config(
    page_title="Smart Word Document Generator",
    page_icon="📄",
    layout="wide"
)
# Gemini API key input
with st.sidebar:
    st.subheader("🔑 Gemini API Configuration")

    gemini_api_key = st.text_input(
        "Enter Gemini API Key",
        type="password",
        placeholder="Paste your Gemini API key"
    )

client = None

if gemini_api_key:
    try:
        client = genai.Client(api_key=gemini_api_key)
        st.sidebar.success("API Key Loaded")
    except Exception as e:
        st.sidebar.error(f"Invalid API Key: {e}")

# Custom CSS for custom premium styling, fonts, and responsiveness
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Outfit', sans-serif;
    }
    
    .gradient-header {
        background: linear-gradient(135deg, #6366F1, #A855F7, #EC4899);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-size: 2.8rem;
        font-weight: 700;
        text-align: center;
        margin-top: 1rem;
        margin-bottom: 0.5rem;
    }
    
    .subtitle {
        font-size: 1.15rem;
        color: #6B7280;
        text-align: center;
        margin-bottom: 2.5rem;
    }
    
    .section-card {
        background-color: var(--secondary-background-color);
        border-radius: 16px;
        padding: 2rem;
        border: 1px solid var(--border-color);
        margin-bottom: 1.5rem;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05), 0 2px 4px -1px rgba(0, 0, 0, 0.03);
    }
    
    .page-title {
        font-size: 1.5rem;
        font-weight: 600;
        margin-bottom: 1rem;
        border-bottom: 2px solid var(--primary-color);
        padding-bottom: 0.5rem;
    }
    
    /* Animation for download button */
    div.stDownloadButton > button {
        background: linear-gradient(135deg, #10B981, #059669) !important;
        color: white !important;
        font-weight: 600 !important;
        border-radius: 8px !important;
        border: none !important;
        padding: 0.75rem 2rem !important;
        transition: transform 0.2s ease, box-shadow 0.2s ease !important;
        width: 100% !important;
        display: inline-block;
        text-align: center;
        box-shadow: 0 4px 14px 0 rgba(16, 185, 129, 0.4) !important;
    }
    
    div.stDownloadButton > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px 0 rgba(16, 185, 129, 0.6) !important;
    }
    
    div.stButton > button {
        border-radius: 8px !important;
        padding: 0.6rem 2rem !important;
        font-weight: 600 !important;
        width: 100% !important;
    }
</style>
""", unsafe_allow_html=True)

# App Title & Subtitle
st.markdown('<div class="gradient-header">Smart Word Document Generator</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">Generate custom formatted Word reports with section-specific layouts, page numbering, and line spacing.</div>', unsafe_allow_html=True)

def generate_body(heading, subheading):

    prompt = f"""
    Act as a professional report writer.

    Heading: {heading}
    Subheading: {subheading}

    Requirements:
    - Professional tone
    - 250-300 words
    - No HTML tags
    - No Markdown
    - No CSS
    - No bullet points
    - Proper paragraph structure
    - Plain text only

    return only the paragraph text
    """

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt
    )

    return response.text

# Session state initialization to hold document buffers
if 'doc_bytes' not in st.session_state:
    st.session_state.doc_bytes = None
if 'generated_filename' not in st.session_state:
    st.session_state.generated_filename = ""

# Sidebar for formatting info
with st.sidebar:
    st.markdown("### 📋 Document Style Rules")
    st.info("""
    The generated document adheres strictly to the following parameters:
    
    * **Times New Roman Font**: Applied to all elements including headers/footers.
    * **First 2 Pages Empty**: Section 1 is blank and border-free.
    * **Section 2 (Page 3+)**: Content starts here, numbering starts at **1**.
    * **Line Spacing**: 1.5 lines document-wide.
    * **Page Margins**: 2.0 Cm (Top, Bottom, Left, Right).
    * **Page Borders**: Single thin border offset from text margins.
    * **Page Numbers**: Outside page border, bottom-right.
    
    **Element Specifications:**
    * **Heading 1**: 16pt, Bold, Centered
    * **Subheading**: 14pt, Bold, Justified
    * **Body Text**: 12pt, Regular, Justified
    """)
    st.markdown("---")
    st.caption("Powered by Streamlit & python-docx")

# Main content splits into 2 columns
col_setup, col_content = st.columns([1, 2], gap="large")

with col_setup:
    st.markdown('<div class="page-title">⚙️ Document Setup</div>', unsafe_allow_html=True)
    
    with st.container(border=True):
        report_title = st.text_input("Report Title", value="Smart Word Report", placeholder="e.g. Project Briefing")
        header_text = st.text_input("Header Text", value="Internal Document", placeholder="e.g. Confidential Info")
        footer_text = st.text_input("Footer Text", value="Smart Generator Corp", placeholder="e.g. Page footer label")
        
        num_pages = st.number_input(
            "Number of Content Pages", 
            min_value=1, 
            max_value=25, 
            value=3, 
            step=1,
            help="This is the number of content pages. The first two pages will remain blank, and content will begin on Page 3."
        )

with col_content:
    st.markdown('<div class="page-title">📝 Page Content Editor</div>', unsafe_allow_html=True)
    
    # Generate dynamic tabs based on page count
    tabs = st.tabs([f"Page {i}" for i in range(1, num_pages + 1)])
    
    pages_input_data = []
    
    for i, tab in enumerate(tabs):
        with tab:
            st.write(f"✍️ **Content for Page {i+1} (starts on Document Page {i+3})**")
            p_heading = st.text_input(f"Heading (Centered, 16pt Bold)", key=f"heading_{i}", value=f"Section {i+1}: Key Findings")
            p_subheading = st.text_input(f"Subheading (Justified, 14pt Bold)", key=f"subheading_{i}", value=f"Overview of Section {i+1}")
            col1, col2 = st.columns([4,1])

            with col1:
                if f"generated_body_{i}" in st.session_state:
                    st.session_state[f"body_input_{i}"] = st.session_state[f"generated_body_{i}"]
                    del st.session_state[f"generated_body_{i}"]

                if f"body_input_{i}" not in st.session_state:
                    st.session_state[f"body_input_{i}"] = ""

                p_body = st.text_area(
                    f"Body Text (Justified, 12pt)",
                    key=f"body_input_{i}",
                    height=200
                )

            with col2:
                st.write("")
                st.write("")

                if st.button(
                    "✨ Generate",
                    key=f"generate_{i}"
                ):

                    if not gemini_api_key:
                        st.error("Please enter your Gemini API key in the sidebar.")
                    elif p_heading and p_subheading:

                        with st.spinner("Generating content..."):

                            generated_text = generate_body(
                                p_heading,
                                p_subheading
                            )

                            st.session_state[f"generated_body_{i}"] = generated_text
                            st.rerun()
            
            pages_input_data.append({
                "heading": p_heading.strip(),
                "subheading": p_subheading.strip(),
                "body": p_body.strip()
            })

    st.markdown("---")
    
    # Generate action buttons
    if st.button("🚀 Compile Document", type="primary"):
        with st.spinner("Generating document..."):
            # Create Document
            doc = Document()
            
            # Configure borders not to surround header/footer
            configure_borders_not_to_surround_header_footer(doc)

            # Configure default style line spacing for the whole document (1.5 line spacing)
            style = doc.styles['Normal']
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            style.paragraph_format.line_spacing = 1.5

            # -------------------------------------------------
            # Section 1: First 2 Pages (Empty, border-free)
            # -------------------------------------------------
            section1 = doc.sections[0]
            section1.top_margin = Cm(2)
            section1.bottom_margin = Cm(2)
            section1.left_margin = Cm(2)
            section1.right_margin = Cm(2)

            # Page 1 empty paragraph
            p1 = doc.add_paragraph()
            p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p1.paragraph_format.line_spacing = 1.5
            doc.add_page_break()

            # Page 2 empty paragraph
            p2 = doc.add_paragraph()
            p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p2.paragraph_format.line_spacing = 1.5

            # -------------------------------------------------
            # Section 2: Page 3 and Onwards (Content pages)
            # -------------------------------------------------
            sec2 = doc.add_section(WD_SECTION_START.NEW_PAGE)
            sec2.header.is_linked_to_previous = False
            sec2.footer.is_linked_to_previous = False

            sec2.top_margin = Cm(2)
            sec2.bottom_margin = Cm(2)
            sec2.left_margin = Cm(2)
            sec2.right_margin = Cm(2)
            
            sec2.header_distance = Cm(0.6)
            sec2.footer_distance = Cm(0.5)

            # Add borders
            add_page_border(sec2)

            # Restart numbering
            restart_page_numbering(sec2)

            # Configure Header
            header = sec2.header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            header_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            header_para.paragraph_format.line_spacing = 1.5
            header_run = header_para.add_run(header_text)
            header_run.font.name = "Times New Roman"
            header_run.font.size = Pt(10)

            # Configure Footer
            footer = sec2.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            footer_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            footer_para.paragraph_format.line_spacing = 1.5
            
            footer_run = footer_para.add_run(footer_text)
            footer_run.font.name = "Times New Roman"
            footer_run.font.size = Pt(10)

            # Add tab stop at the right margin edge to align page number to the right
            usable_width = sec2.page_width - sec2.left_margin - sec2.right_margin
            footer_para.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
            
            # Add tab space to move cursor to the right
            footer_para.add_run("\t\t")
            
            # Add page number
            add_page_number(footer_para)

            # Center Report Title at the top of Section 2
            if report_title.strip():
                title_para = doc.add_paragraph()
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                title_para.paragraph_format.line_spacing = 1.5
                title_run = title_para.add_run(report_title)
                title_run.bold = True
                title_run.font.name = "Times New Roman"
                title_run.font.size = Pt(18)
                
                spacer = doc.add_paragraph()
                spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                spacer.paragraph_format.line_spacing = 1.5

            # Write dynamically entered page contents
            for idx, page in enumerate(pages_input_data):
                if page["heading"]:
                    add_main_heading(doc, page["heading"])
                if page["subheading"]:
                    add_sub_heading(doc, page["subheading"])
                if page["body"]:
                    add_body_text(doc, page["body"])

                if idx < len(pages_input_data) - 1:
                    doc.add_page_break()

            # Save file to a BytesIO object for Streamlit downloads
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            # Store in session state
            st.session_state.doc_bytes = doc_buffer.getvalue()
            st.session_state.generated_filename = report_title.replace(" ", "_") + ".docx"
            
            st.success("🎉 Word document compiled successfully!")

    # Display download button if bytes exist in session state
    if st.session_state.doc_bytes is not None:
        st.markdown("<br>", unsafe_allow_html=True)
        st.download_button(
            label="📥 Download Generated Word Document (.docx)",
            data=st.session_state.doc_bytes,
            file_name=st.session_state.generated_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
